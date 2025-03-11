#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word to Blocky Formatter

Tento skript extrahuje údaje z Word dokumentu (.docx) a vytvorí tabuľku 
v rovnakom formáte ako vo Word dokumente. Výsledok je možné použiť 
v aplikácii Blocky.

Použitie:
    python word_to_blocky_formatter.py input.docx output.txt

Požiadavky:
    - Python 3.6+
    - python-docx

Inštalácia požiadaviek:
    pip install python-docx
"""

import sys
import os
import re
from datetime import datetime
from docx import Document


def extract_table_format_from_word(input_file):
    """
    Extrahuje údaje a formátovanie tabuľky z Word dokumentu.
    
    Args:
        input_file (str): Cesta k vstupnému Word dokumentu
    
    Returns:
        dict: Slovník obsahujúci údaje a formátovanie tabuľky
    """
    try:
        # Načítanie Word dokumentu
        print(f"Načítavam Word dokument: {input_file}")
        doc = Document(input_file)
        
        table_data = {
            "headers": [],
            "rows": [],
            "title": "",
            "has_borders": False,
            "column_widths": [],
            "has_header_formatting": False
        }
        
        # Hľadanie nadpisu tabuľky
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and not table_data["title"]:
                table_data["title"] = paragraph.text.strip()
                break
        
        # Pokus o extrakciu údajov z tabuliek
        if len(doc.tables) > 0:
            print(f"Dokument obsahuje {len(doc.tables)} tabuliek.")
            
            # Berieme prvú tabuľku
            table = doc.tables[0]
            print(f"Spracovávam tabuľku...")
            
            # Zistenie, či tabuľka má ohraničenie
            if hasattr(table, 'style') and table.style:
                table_data["has_borders"] = "Table Grid" in table.style.name or "Grid" in table.style.name
            
            # Extrakcia hlavičky tabuľky
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    table_data["headers"].append(cell.text.strip())
                
                # Zistenie, či hlavička má špeciálne formátovanie
                table_data["has_header_formatting"] = True  # Predpokladáme, že má
                
                # Extrakcia šírky stĺpcov (približne)
                for cell in header_row.cells:
                    if hasattr(cell, 'width'):
                        table_data["column_widths"].append(cell.width)
                    else:
                        table_data["column_widths"].append(None)
            
            # Extrakcia riadkov tabuľky
            for row_index, row in enumerate(table.rows):
                if row_index == 0:  # Preskočenie hlavičky
                    continue
                
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                
                table_data["rows"].append(row_data)
        
        print(f"Extrakcia dokončená. Nájdených {len(table_data['rows'])} riadkov.")
        return table_data
    
    except Exception as e:
        print(f"Chyba pri extrakcii údajov z Word dokumentu: {str(e)}")
        return None


def create_formatted_table(table_data, output_file, input_file_name):
    """
    Vytvorí formátovanú tabuľku na základe extrahovaných údajov.
    
    Args:
        table_data (dict): Slovník obsahujúci údaje a formátovanie tabuľky
        output_file (str): Cesta k výstupnému textovému súboru
        input_file_name (str): Názov vstupného súboru
    
    Returns:
        bool: True, ak sa podarilo vytvoriť tabuľku, inak False
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            # Hlavička dokumentu
            f.write(f"# Tabuľka vytvorená z dokumentu: {input_file_name}\n")
            f.write(f"# Vytvorené: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
            f.write("#\n\n")
            
            # Nadpis tabuľky
            if table_data["title"]:
                f.write(f"{table_data['title']}\n\n")
            
            # Určenie šírky stĺpcov
            column_widths = []
            for i, header in enumerate(table_data["headers"]):
                # Nájdenie maximálnej dĺžky textu v stĺpci
                max_width = len(header)
                for row in table_data["rows"]:
                    if i < len(row):
                        max_width = max(max_width, len(row[i]))
                
                # Pridanie rezervy
                column_widths.append(max_width + 2)
            
            # Vytvorenie horného ohraničenia tabuľky
            if table_data["has_borders"]:
                f.write("+" + "+".join("-" * width for width in column_widths) + "+\n")
            
            # Vytvorenie hlavičky tabuľky
            header_row = "|"
            for i, header in enumerate(table_data["headers"]):
                width = column_widths[i]
                header_row += f" {header.ljust(width - 2)} |"
            f.write(header_row + "\n")
            
            # Vytvorenie oddeľovača hlavičky
            if table_data["has_borders"] or table_data["has_header_formatting"]:
                f.write("+" + "+".join("-" * width for width in column_widths) + "+\n")
            
            # Vytvorenie riadkov tabuľky
            for row in table_data["rows"]:
                row_text = "|"
                for i, cell in enumerate(row):
                    if i < len(column_widths):
                        width = column_widths[i]
                        row_text += f" {cell.ljust(width - 2)} |"
                f.write(row_text + "\n")
            
            # Vytvorenie spodného ohraničenia tabuľky
            if table_data["has_borders"]:
                f.write("+" + "+".join("-" * width for width in column_widths) + "+\n")
            
            # Výpočet celkovej sumy (ak je to relevantné)
            try:
                # Predpokladáme, že posledný stĺpec obsahuje sumy
                total_sum = 0
                for row in table_data["rows"]:
                    if len(row) > 0:
                        last_cell = row[-1]
                        # Extrakcia čísla zo sumy
                        amount_match = re.search(r'([\-]?\d+[.,]?\d*)', last_cell)
                        if amount_match:
                            amount_str = amount_match.group(1).replace(',', '.')
                            total_sum += float(amount_str)
                
                # Pridanie celkovej sumy pod tabuľku
                f.write("\n")
                f.write(f"Celková suma: {total_sum:.2f} €\n")
            except:
                # Ak sa nepodarí vypočítať sumu, ignorujeme to
                pass
            
            # Pridanie informácií pre import do aplikácie Blocky
            f.write("\n")
            f.write("# Informácie pre import do aplikácie Blocky:\n")
            f.write("# Pre každý riadok tabuľky (okrem hlavičky) vytvorte položku v aplikácii:\n")
            
            for row in table_data["rows"]:
                if len(row) >= 2:  # Predpokladáme, že prvý stĺpec je popis a posledný je suma
                    description = row[0]
                    amount_cell = row[-1]
                    
                    # Extrakcia čísla zo sumy
                    amount_match = re.search(r'([\-]?\d+[.,]?\d*)', amount_cell)
                    if amount_match:
                        amount_str = amount_match.group(1).replace(',', '.')
                        amount = float(amount_str)
                        
                        f.write(f"# - Popis: {description}, Suma: {amount:.2f} €\n")
        
        print(f"Tabuľka bola úspešne vytvorená. Výstupný súbor: {output_file}")
        return True
    
    except Exception as e:
        print(f"Chyba pri vytváraní tabuľky: {str(e)}")
        return False


def print_usage():
    """Zobrazí návod na použitie"""
    print("Použitie:")
    print("  python word_to_blocky_formatter.py input.docx output.txt")
    print()
    print("Parametre:")
    print("  input.docx  - Vstupný Word dokument")
    print("  output.txt  - Výstupný textový súbor")


def main():
    """Hlavná funkcia"""
    # Kontrola argumentov
    if len(sys.argv) != 3:
        print_usage()
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Kontrola, či vstupný súbor existuje
    if not os.path.isfile(input_file):
        print(f"Chyba: Vstupný súbor '{input_file}' neexistuje")
        return
    
    # Kontrola, či vstupný súbor je Word dokument
    if not input_file.endswith(('.docx')):
        print(f"Upozornenie: Vstupný súbor '{input_file}' nemusí byť Word dokument")
    
    # Extrakcia údajov a formátovania tabuľky z Word dokumentu
    table_data = extract_table_format_from_word(input_file)
    
    if not table_data:
        print("Neboli nájdené žiadne údaje v tabuľke.")
        return
    
    # Vytvorenie formátovanej tabuľky
    create_formatted_table(table_data, output_file, os.path.basename(input_file))


if __name__ == "__main__":
    main() 