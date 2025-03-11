#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word to Blocky Converter

Tento skript extrahuje údaje z Word dokumentu (.docx) a pripraví ich 
na import do aplikácie Blocky. Skript vytvorí textový súbor s položkami 
vo formáte, ktorý je možné manuálne zadať do aplikácie.

Použitie:
    python word_to_blocky.py input.docx output.txt

Požiadavky:
    - Python 3.6+
    - python-docx

Inštalácia požiadaviek:
    pip install python-docx
"""

import sys
import os
import re
from datetime import datetime
from docx import Document


def extract_data_from_word(input_file):
    """
    Extrahuje údaje z Word dokumentu.
    
    Args:
        input_file (str): Cesta k vstupnému Word dokumentu
    
    Returns:
        list: Zoznam položiek vo formáte [(popis, suma)]
    """
    try:
        # Načítanie Word dokumentu
        print(f"Načítavam Word dokument: {input_file}")
        doc = Document(input_file)
        
        items = []
        
        # Pokus o extrakciu údajov z tabuliek
        if len(doc.tables) > 0:
            print(f"Dokument obsahuje {len(doc.tables)} tabuliek.")
            for table_index, table in enumerate(doc.tables):
                print(f"Spracovávam tabuľku {table_index + 1}...")
                
                # Predpokladáme, že prvý stĺpec je popis a druhý je suma
                for row_index, row in enumerate(table.rows):
                    # Preskočenie hlavičky tabuľky
                    if row_index == 0:
                        continue
                    
                    if len(row.cells) >= 2:
                        description = row.cells[0].text.strip()
                        amount_text = row.cells[1].text.strip()
                        
                        # Pokus o extrakciu čísla zo sumy
                        amount = extract_amount(amount_text)
                        
                        if amount is not None:
                            items.append((description, amount))
        
        # Ak neboli nájdené žiadne tabuľky alebo údaje v tabuľkách, 
        # pokúsime sa extrahovať údaje z textu
        if len(items) == 0:
            print("Neboli nájdené žiadne údaje v tabuľkách, pokúšam sa extrahovať údaje z textu...")
            
            # Hľadanie vzoru "popis: suma" alebo "popis suma €"
            pattern = r'([^:]+):\s*([\-]?\d+[.,]?\d*)\s*€?'
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Hľadanie vzoru v texte
                    matches = re.findall(pattern, text)
                    for match in matches:
                        description = match[0].strip()
                        amount = extract_amount(match[1])
                        if amount is not None:
                            items.append((description, amount))
        
        print(f"Celkovo nájdených položiek: {len(items)}")
        return items
    
    except Exception as e:
        print(f"Chyba pri extrakcii údajov z Word dokumentu: {str(e)}")
        return []


def extract_amount(text):
    """
    Extrahuje číselnú hodnotu zo zadaného textu.
    
    Args:
        text (str): Text obsahujúci číselnú hodnotu
    
    Returns:
        float: Extrahovaná číselná hodnota alebo None, ak sa nepodarilo extrahovať
    """
    try:
        # Odstránenie nečíselných znakov okrem mínus, bodky a čiarky
        cleaned_text = re.sub(r'[^\d\-.,]', '', text)
        
        # Nahradenie čiarky bodkou
        cleaned_text = cleaned_text.replace(',', '.')
        
        # Konverzia na číslo
        return float(cleaned_text)
    except (ValueError, TypeError):
        return None


def save_to_text_file(items, output_file, input_file_name):
    """
    Uloží extrahované položky do textového súboru.
    
    Args:
        items (list): Zoznam položiek vo formáte [(popis, suma)]
        output_file (str): Cesta k výstupnému textovému súboru
        input_file_name (str): Názov vstupného súboru
    
    Returns:
        bool: True, ak sa podarilo uložiť súbor, inak False
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# Položky pre import do aplikácie Blocky\n")
            f.write("# Vytvorené: " + datetime.now().strftime("%d.%m.%Y %H:%M:%S") + "\n")
            f.write("# Zdrojový súbor: " + input_file_name + "\n")
            f.write("#\n")
            f.write("# Formát: Popis | Suma\n")
            f.write("# Záporné sumy sú označené znamienkom mínus (-)\n")
            f.write("#\n")
            
            # Zápis položiek
            total_amount = 0
            for description, amount in items:
                total_amount += amount
                
                # Formátovanie sumy
                amount_str = f"{amount:.2f}".replace('.', ',')
                
                # Zápis položky
                f.write(f"{description} | {amount_str}\n")
            
            # Zápis súhrnu
            f.write("#\n")
            f.write(f"# Celková suma: {total_amount:.2f}\n")
            f.write(f"# Počet položiek: {len(items)}\n")
        
        print(f"Konverzia dokončená. Výstupný súbor: {output_file}")
        print(f"Celkový počet položiek: {len(items)}")
        print(f"Celková suma: {total_amount:.2f}")
        
        return True
    
    except Exception as e:
        print(f"Chyba pri ukladaní do textového súboru: {str(e)}")
        return False


def print_usage():
    """Zobrazí návod na použitie"""
    print("Použitie:")
    print("  python word_to_blocky.py input.docx output.txt")
    print()
    print("Parametre:")
    print("  input.docx  - Vstupný Word dokument")
    print("  output.txt  - Výstupný textový súbor")


def main():
    """Hlavná funkcia"""
    # Kontrola argumentov
    if len(sys.argv) != 3:
        print_usage()
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Kontrola, či vstupný súbor existuje
    if not os.path.isfile(input_file):
        print(f"Chyba: Vstupný súbor '{input_file}' neexistuje")
        return
    
    # Kontrola, či vstupný súbor je Word dokument
    if not input_file.endswith(('.docx')):
        print(f"Upozornenie: Vstupný súbor '{input_file}' nemusí byť Word dokument")
    
    # Extrakcia údajov z Word dokumentu
    items = extract_data_from_word(input_file)
    
    if len(items) == 0:
        print("Neboli nájdené žiadne položky na import.")
        return
    
    # Uloženie do textového súboru
    save_to_text_file(items, output_file, os.path.basename(input_file))


if __name__ == "__main__":
    main() 